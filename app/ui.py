import re
import datetime
import time
from io import BytesIO

import streamlit as st
from docx import Document
from loader import load_knowledge, load_prompt
from generate import stream_proposal, stream_answers, extract_job_details
from config import KNOWLEDGE_PATH, PROMPT_PATH
from crm import log_to_crm

st.set_page_config(page_title="Upwork AI Agent", layout="wide")

@st.cache_resource
def get_knowledge():
    return load_knowledge(KNOWLEDGE_PATH)

@st.cache_resource
def get_prompt_style():
    return load_prompt(PROMPT_PATH)

knowledge    = get_knowledge()
prompt_style = get_prompt_style()

if "proposal" not in st.session_state:
    st.session_state.proposal = ""
if "answers" not in st.session_state:
    st.session_state.answers = ""
if "job_details" not in st.session_state:
    st.session_state.job_details = {}

st.title("📊 Upwork Proposal Generator")
st.caption("Powered by Groq")

col1, col2 = st.columns([1, 1], gap="large")

with col1:
    st.subheader("📋 Job Description")
    job_description = st.text_area(
        label="Job description",
        height=250,
        placeholder="Paste the full Upwork job page here...",
        label_visibility="collapsed",
        key="job_input"
    )

    st.subheader("🔗 Job Link")
    job_link = st.text_input(
        label="Job Link",
        placeholder="https://www.upwork.com/jobs/...",
        label_visibility="collapsed",
        key="job_link_input"
    )

    st.subheader("❓ Client Questions (optional)")
    client_questions = st.text_area(
        label="Client Questions",
        height=150,
        placeholder="Paste any questions the client asked...",
        label_visibility="collapsed",
        key="questions_input"
    )

    generate_btn   = st.button("✨ Generate Proposal", use_container_width=True, type="primary")
    regenerate_btn = st.button("🔄 Regenerate", use_container_width=True)


def clean_proposal(text: str) -> str:
    sentences = re.split(r'(?<=[.!?])\s+', text)
    return " ".join([s for s in sentences if not s.strip().startswith("Most")])


def show_copyable(text: str):
    st.caption("👆 Hover over the box and click the clipboard icon to copy")
    st.code(text, language=None)


with col2:
    if generate_btn or regenerate_btn:
        if not job_description.strip():
            st.warning("Please paste a job description first.")
        else:
            st.session_state.proposal = ""
            st.session_state.answers = ""

            # ── Extract job details ──────────────────────────────────
            with st.spinner("Extracting job details..."):
                details = extract_job_details(job_description)
                st.session_state.job_details = details

            # ── Show extracted details ───────────────────────────────
            with st.expander("📌 Extracted Job Details", expanded=True):
                st.write(f"**Title:** {details.get('title', 'N/A')}")
                st.write(f"**Budget:** {details.get('budget', 'N/A')}")
                st.write(f"**Duration:** {details.get('duration', 'N/A')}")
                st.write(f"**Connects:** {details.get('connects', 'N/A')}")

            # ── Stream Proposal ──────────────────────────────────────
            st.subheader("✍️ Proposal")
            proposal_box = st.empty()
            full_proposal = ""

            for token in stream_proposal(job_description, client_questions, knowledge, prompt_style):
                full_proposal += token
                proposal_box.markdown(full_proposal + "▌")

            full_proposal = clean_proposal(full_proposal)
            st.session_state.proposal = full_proposal
            proposal_box.markdown(full_proposal)
            show_copyable(full_proposal)

            # ── Auto log to CRM ──────────────────────────────────────
            result = log_to_crm(
                date=str(datetime.date.today()),
                job_title=details.get("title", "Unknown"),
                job_link=job_link.strip() if job_link.strip() else "Not provided",
                budget=details.get("budget", "Unknown"),
                duration=details.get("duration", "Unknown"),
                connects=details.get("connects", "Unknown"),
                proposal=full_proposal
            )
            if result is True:
                st.success("✅ Logged to CRM")
            else:
                st.warning(f"CRM log failed: {result}")

            # ── Stream Answers (only if questions provided) ──────────
            if client_questions.strip():
                time.sleep(15)
                st.divider()
                st.subheader("💬 Answers to Client Questions")
                answers_box = st.empty()
                full_answers = ""
                try:
                    for token in stream_answers(job_description, client_questions, knowledge, prompt_style):
                        full_answers += token
                        answers_box.markdown(full_answers + "▌")
                    answers_box.markdown(full_answers)
                    st.session_state.answers = full_answers
                    show_copyable(full_answers)
                except Exception as e:
                    st.error(f"Answers failed: {e}")

    elif st.session_state.proposal:
        if st.session_state.job_details:
            with st.expander("📌 Extracted Job Details"):
                d = st.session_state.job_details
                st.write(f"**Title:** {d.get('title', 'N/A')}")
                st.write(f"**Budget:** {d.get('budget', 'N/A')}")
                st.write(f"**Duration:** {d.get('duration', 'N/A')}")
                st.write(f"**Connects:** {d.get('connects', 'N/A')}")

        st.subheader("✍️ Proposal")
        st.markdown(st.session_state.proposal)
        show_copyable(st.session_state.proposal)

        if st.session_state.answers:
            st.divider()
            st.subheader("💬 Answers to Client Questions")
            st.markdown(st.session_state.answers)
            show_copyable(st.session_state.answers)

    else:
        st.info("Your proposal will appear here once generated.")

    # ── Save to Knowledge Base ─────────────────────────────────────── (this is the old ui)
    # if st.session_state.proposal:
    #     st.divider()
    #     if st.button("✅ Save to Knowledge Base"):
    #         doc = Document()
    #         doc.add_paragraph(f"DATE: {datetime.date.today()}")
    #         doc.add_paragraph("")
    #         doc.add_paragraph("PROPOSAL:")
    #         doc.add_paragraph(st.session_state.proposal)
    #         if st.session_state.answers:
    #             doc.add_paragraph("")
    #             doc.add_paragraph("QUESTION ANSWERS:")
    #             doc.add_paragraph(st.session_state.answers)

    #         buffer = BytesIO()
    #         doc.save(buffer)
    #         buffer.seek(0)

    #         st.download_button(
    #             label="📥 Download .docx",
    #             data=buffer,
    #             file_name=f"proposal_{datetime.date.today()}.docx",
    #             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    #         )
    if st.session_state.proposal:
        st.divider()
    
    # Clear All button
    if st.button("🗑️ Clear All", use_container_width=True):
        st.session_state.proposal = ""
        st.session_state.answers = ""
        st.session_state.job_details = {}
        st.rerun()

    # Auto download on click
    doc = Document()
    doc.add_paragraph(f"DATE: {datetime.date.today()}")
    doc.add_paragraph("")
    doc.add_paragraph("PROPOSAL:")
    doc.add_paragraph(st.session_state.proposal)
    if st.session_state.answers:
        doc.add_paragraph("")
        doc.add_paragraph("QUESTION ANSWERS:")
        doc.add_paragraph(st.session_state.answers)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="✅ Save to Knowledge Base",
        data=buffer,
        file_name=f"proposal_{datetime.date.today()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )